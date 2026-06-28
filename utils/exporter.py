"""
utils/exporter.py - Export Engine
AI Study Notes Generator
-----------------------------------------
Handles exporting generated study content as:
- PDF  (via reportlab)
- DOCX (via python-docx)
- TXT  (plain text)
"""

import os
import re
import logging
from datetime import datetime

# PDF generation
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer,
    HRFlowable, PageBreak, Table, TableStyle
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER

# DOCX generation
# pyrefly: ignore [missing-import]
from docx import Document
# pyrefly: ignore [missing-import]
from docx.shared import Inches, Pt, RGBColor
# pyrefly: ignore [missing-import]
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────

def _sanitize_filename(name: str) -> str:
    """Remove invalid characters from filename."""
    return re.sub(r'[^\w\-_.]', '_', name)


def _get_timestamp() -> str:
    """Return current timestamp string."""
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _strip_markdown(text: str) -> str:
    """
    Strip basic markdown formatting for plain text export.
    Keeps structure but removes symbols.
    """
    # Remove bold/italic markers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'__(.+?)__',     r'\1', text)

    # Convert headers to uppercase labels
    text = re.sub(r'^#{1,2}\s+(.+)$', r'\n\1\n' + '─' * 40, text, flags=re.MULTILINE)
    text = re.sub(r'^#{3,}\s+(.+)$',  r'\1:', text, flags=re.MULTILINE)

    return text.strip()


# ─────────────────────────────────────────
# TXT Export
# ─────────────────────────────────────────

def export_txt(content: dict, output_dir: str, filename: str = None) -> str:
    """
    Export all generated content as a TXT file.

    Args:
        content:    Dict of generated content sections
        output_dir: Directory to save the file
        filename:   Optional custom filename

    Returns:
        Full path to the saved TXT file
    """
    os.makedirs(output_dir, exist_ok=True)

    if not filename:
        filename = f"study_notes_{_get_timestamp()}.txt"

    filepath = os.path.join(output_dir, filename)

    lines = []
    divider = "=" * 60

    lines.append("AI STUDY NOTES GENERATOR")
    lines.append(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    lines.append(divider)

    # ── Summary ──
    if "summary" in content:
        lines.append("\n📋 SUMMARY")
        lines.append(divider)
        summary = content["summary"]
        for stype in ["short", "medium", "detailed"]:
            if stype in summary:
                lines.append(f"\n[{stype.upper()} SUMMARY]")
                lines.append(_strip_markdown(summary[stype].get("content", "")))

    # ── Study Notes ──
    if "study_notes" in content:
        lines.append(f"\n\n📚 STUDY NOTES")
        lines.append(divider)
        lines.append(_strip_markdown(content["study_notes"].get("content", "")))

    # ── Flashcards ──
    if "flashcards" in content:
        lines.append(f"\n\n🃏 FLASHCARDS")
        lines.append(divider)
        cards = content["flashcards"].get("flashcards", [])
        for card in cards:
            lines.append(f"\nQ{card.get('id', '')}: {card.get('question', '')}")
            lines.append(f"A:  {card.get('answer', '')}")
            lines.append("-" * 40)

    # ── Quiz ──
    if "quiz" in content:
        quiz = content["quiz"]
        lines.append(f"\n\n📝 QUIZ")
        lines.append(divider)

        if "mcq" in quiz:
            lines.append("\n[MULTIPLE CHOICE]")
            for q in quiz["mcq"]:
                lines.append(f"\n{q.get('id', '')}. {q.get('question', '')}")
                for opt, val in q.get("options", {}).items():
                    lines.append(f"   {opt}) {val}")
                lines.append(f"   ✅ Answer: {q.get('correct_answer', '')}")
                lines.append(f"   💡 {q.get('explanation', '')}")

        if "true_false" in quiz:
            lines.append("\n[TRUE / FALSE]")
            for q in quiz["true_false"]:
                ans = "True" if q.get("correct_answer") else "False"
                lines.append(f"\n{q.get('id', '')}. {q.get('question', '')}")
                lines.append(f"   ✅ Answer: {ans}")

        if "fill_blank" in quiz:
            lines.append("\n[FILL IN THE BLANK]")
            for q in quiz["fill_blank"]:
                lines.append(f"\n{q.get('id', '')}. {q.get('question', '')}")
                lines.append(f"   ✅ Answer: {q.get('answer', '')}")

        if "short_answer" in quiz:
            lines.append("\n[SHORT ANSWER]")
            for q in quiz["short_answer"]:
                lines.append(f"\n{q.get('id', '')}. {q.get('question', '')}")
                lines.append(f"   ✅ Answer: {q.get('answer', '')}")

    # ── Important Questions ──
    if "important_questions" in content:
        iq = content["important_questions"]
        lines.append(f"\n\n🎯 IMPORTANT EXAM QUESTIONS")
        lines.append(divider)
        for level in ["easy", "medium", "hard"]:
            if level in iq:
                lines.append(f"\n[{level.upper()}]")
                for q in iq[level]:
                    lines.append(f"\n{q.get('id', '')}. {q.get('question', '')} [{q.get('marks', '')} marks]")
                    lines.append(f"   💡 Hint: {q.get('hint', '')}")

    # ── Cheat Sheet ──
    if "cheat_sheet" in content:
        lines.append(f"\n\n⚡ CHEAT SHEET")
        lines.append(divider)
        lines.append(_strip_markdown(content["cheat_sheet"].get("content", "")))

    with open(filepath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    logger.info(f"✅ TXT exported: {filepath}")
    return filepath


# ─────────────────────────────────────────
# DOCX Export
# ─────────────────────────────────────────

def export_docx(content: dict, output_dir: str, filename: str = None) -> str:
    """
    Export all generated content as a DOCX file.

    Args:
        content:    Dict of generated content sections
        output_dir: Directory to save the file
        filename:   Optional custom filename

    Returns:
        Full path to the saved DOCX file
    """
    os.makedirs(output_dir, exist_ok=True)

    if not filename:
        filename = f"study_notes_{_get_timestamp()}.docx"

    filepath = os.path.join(output_dir, filename)
    doc = Document()

    # ── Document Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── Title Page ──
    title = doc.add_heading('AI Study Notes', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        f"Generated on {datetime.now().strftime('%B %d, %Y at %H:%M')}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ── Summary Section ──
    if "summary" in content:
        doc.add_heading('📋 Summary', level=1)
        summary = content["summary"]

        for stype in ["short", "medium", "detailed"]:
            if stype in summary:
                doc.add_heading(f'{stype.capitalize()} Summary', level=2)
                text = summary[stype].get("content", "")
                doc.add_paragraph(_strip_markdown(text))
                doc.add_paragraph()

        doc.add_page_break()

    # ── Study Notes Section ──
    if "study_notes" in content:
        doc.add_heading('📚 Study Notes', level=1)
        notes_text = content["study_notes"].get("content", "")

        for line in notes_text.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line[2:])
            else:
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                doc.add_paragraph(clean)

        doc.add_page_break()

    # ── Flashcards Section ──
    if "flashcards" in content:
        doc.add_heading('🃏 Flashcards', level=1)
        cards = content["flashcards"].get("flashcards", [])

        for card in cards:
            p = doc.add_paragraph()
            p.add_run(f"Q{card.get('id', '')}: ").bold = True
            p.add_run(card.get('question', ''))

            p2 = doc.add_paragraph()
            r = p2.add_run("Answer: ")
            r.bold = True
            r.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
            p2.add_run(card.get('answer', ''))
            doc.add_paragraph('─' * 50)

        doc.add_page_break()

    # ── Cheat Sheet Section ──
    if "cheat_sheet" in content:
        doc.add_heading('⚡ Cheat Sheet', level=1)
        cheat_text = content["cheat_sheet"].get("content", "")

        for line in cheat_text.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line[2:])
            else:
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                doc.add_paragraph(clean)

    doc.save(filepath)
    logger.info(f"✅ DOCX exported: {filepath}")
    return filepath


# ─────────────────────────────────────────
# PDF Export
# ─────────────────────────────────────────

def export_pdf(content: dict, output_dir: str, filename: str = None) -> str:
    """
    Export all generated content as a PDF file.

    Args:
        content:    Dict of generated content sections
        output_dir: Directory to save the file
        filename:   Optional custom filename

    Returns:
        Full path to the saved PDF file
    """
    os.makedirs(output_dir, exist_ok=True)

    if not filename:
        filename = f"study_notes_{_get_timestamp()}.pdf"

    filepath = os.path.join(output_dir, filename)

    # ── Page Setup ──
    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        rightMargin=inch * 0.75,
        leftMargin=inch * 0.75,
        topMargin=inch,
        bottomMargin=inch
    )

    # ── Styles ──
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=24,
        textColor=colors.HexColor('#4F46E5'),
        spaceAfter=12,
        alignment=TA_CENTER
    )

    h1_style = ParagraphStyle(
        'CustomH1',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#4F46E5'),
        spaceBefore=16,
        spaceAfter=8,
        borderPad=4
    )

    h2_style = ParagraphStyle(
        'CustomH2',
        parent=styles['Heading2'],
        fontSize=13,
        textColor=colors.HexColor('#10B981'),
        spaceBefore=12,
        spaceAfter=6
    )

    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        leading=16,
        spaceAfter=6
    )

    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        fontSize=10,
        leading=16,
        leftIndent=20,
        spaceAfter=4,
        bulletIndent=10
    )

    story = []

    # ── Title ──
    story.append(Paragraph("📚 AI Study Notes", title_style))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}",
        styles['Normal']
    ))
    story.append(Spacer(1, 0.3 * inch))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#4F46E5')))
    story.append(Spacer(1, 0.2 * inch))

    def add_section_content(text: str):
        """Parse and add markdown-style text to PDF story."""
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.1 * inch))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], h1_style))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], h2_style))
            elif line.startswith('- ') or line.startswith('* '):
                clean = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line[2:])
                story.append(Paragraph(f"• {clean}", bullet_style))
            else:
                clean = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
                story.append(Paragraph(clean, body_style))

    # ── Summary ──
    if "summary" in content:
        story.append(Paragraph("📋 Summary", h1_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E2E8F0')))
        summary = content["summary"]
        for stype in ["short", "medium", "detailed"]:
            if stype in summary:
                story.append(Paragraph(f"{stype.capitalize()} Summary", h2_style))
                add_section_content(summary[stype].get("content", ""))
        story.append(PageBreak())

    # ── Study Notes ──
    if "study_notes" in content:
        story.append(Paragraph("📚 Study Notes", h1_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E2E8F0')))
        add_section_content(content["study_notes"].get("content", ""))
        story.append(PageBreak())

    # ── Flashcards ──
    if "flashcards" in content:
        story.append(Paragraph("🃏 Flashcards", h1_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E2E8F0')))
        cards = content["flashcards"].get("flashcards", [])

        for card in cards:
            story.append(Paragraph(
                f"<b>Q{card.get('id', '')}:</b> {card.get('question', '')}",
                body_style
            ))
            story.append(Paragraph(
                f"<b>Answer:</b> {card.get('answer', '')}",
                ParagraphStyle('Answer', parent=body_style,
                               textColor=colors.HexColor('#10B981'), leftIndent=15)
            ))
            story.append(HRFlowable(width="100%", thickness=0.5,
                                     color=colors.HexColor('#E2E8F0')))
            story.append(Spacer(1, 0.05 * inch))

        story.append(PageBreak())

    # ── Cheat Sheet ──
    if "cheat_sheet" in content:
        story.append(Paragraph("⚡ Cheat Sheet", h1_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E2E8F0')))
        add_section_content(content["cheat_sheet"].get("content", ""))

    # ── Build PDF ──
    doc.build(story)
    logger.info(f"✅ PDF exported: {filepath}")
    return filepath


# ─────────────────────────────────────────
# Master Export Function
# ─────────────────────────────────────────

def export_content(content: dict, export_format: str, output_dir: str, filename: str = None) -> str:
    """
    Master export function — routes to correct exporter.

    Args:
        content:       Generated content dict
        export_format: 'pdf', 'docx', or 'txt'
        output_dir:    Output directory path
        filename:      Optional custom filename

    Returns:
        Full path to exported file
    """
    exporters = {
        "pdf":  export_pdf,
        "docx": export_docx,
        "txt":  export_txt,
    }

    fmt = export_format.lower().strip(".")

    if fmt not in exporters:
        raise ValueError(
            f"Unsupported export format: {export_format}. "
            f"Choose from: {', '.join(exporters.keys())}"
        )

    logger.info(f"📤 Exporting as {fmt.upper()}...")
    return exporters[fmt](content, output_dir, filename)
